import io
import docx
import pandas as pd
import pytesseract
from pdf2image import convert_from_bytes
import pymupdf as fitz
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="PDF Studio API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def health_check():
    return {"status": "ok", "service": "PDF Studio Backend"}

@app.post("/api/ocr")
async def process_ocr(file: UploadFile = File(...), lang: str = Form("por")):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Apenas arquivos PDF são aceitos.")

    pdf_bytes = await file.read()
    try:
        images = convert_from_bytes(pdf_bytes, dpi=120)
        if not images:
            raise HTTPException(status_code=400, detail="O PDF está vazio.")
        if len(images) > 10:
            raise HTTPException(status_code=400, detail="Envie um PDF de até 10 páginas.")

        pdf_pesquisavel = fitz.open()
        for img in images:
            pdf_page_bytes = pytesseract.image_to_pdf_or_hocr(img, lang=lang, extension='pdf')
            page_doc = fitz.open("pdf", pdf_page_bytes)
            pdf_pesquisavel.insert_pdf(page_doc)
            page_doc.close()

        output_pdf_bytes = pdf_pesquisavel.tobytes()
        pdf_pesquisavel.close()

        return Response(
            content=output_pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=ocr_{file.filename}"}
        )
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro no OCR: {str(e)}")

@app.post("/api/convert/word")
async def convert_to_word(file: UploadFile = File(...), lang: str = Form("por")):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Apenas arquivos PDF são aceitos.")

    pdf_bytes = await file.read()
    try:
        images = convert_from_bytes(pdf_bytes, dpi=120)
        if len(images) > 10:
            raise HTTPException(status_code=400, detail="Envie um PDF de até 10 páginas.")

        doc = docx.Document()
        for idx, img in enumerate(images):
            text = pytesseract.image_to_string(img, lang=lang)
            doc.add_heading(f"Página {idx + 1}", level=2)
            doc.add_paragraph(text)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={file.filename}.docx"}
        )
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao converter para Word: {str(e)}")

@app.post("/api/convert/excel")
async def convert_to_excel(file: UploadFile = File(...), lang: str = Form("por")):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Apenas arquivos PDF são aceitos.")

    pdf_bytes = await file.read()
    try:
        images = convert_from_bytes(pdf_bytes, dpi=120)
        if len(images) > 10:
            raise HTTPException(status_code=400, detail="Envie um PDF de até 10 páginas.")

        data = []
        for idx, img in enumerate(images):
            text = pytesseract.image_to_string(img, lang=lang)
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            for line in lines:
                data.append({"Página": idx + 1, "Texto": line})

        df = pd.DataFrame(data)
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name="OCR")

        buffer.seek(0)

        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={file.filename}.xlsx"}
        )
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao converter para Excel: {str(e)}")
